import json
import pandas as pd
from docx import Document
from fpdf import FPDF
import os
from openpyxl import load_workbook

json_path = r'C:\Users\Bokhan\Desktop\ИРНИТУ\Правоведение\dannie.json'

output_dir = r'C:\Users\Bokhan\Desktop\lol'
os.makedirs(output_dir, exist_ok=True) 

with open(json_path, 'r', encoding='utf-8') as f:
    data = json.load(f)

df = pd.DataFrame(data)
df["cost"] = df['quantity'] * df['price']

excel_path = os.path.join(output_dir, 'report.xlsx')
df.to_excel(excel_path, index=False, sheet_name='Goods')
print("Excel создан")
wb = load_workbook(excel_path)
ws = wb['Goods']
ws.cell(row=ws.max_row + 1, column=1, value="is:")
sum_f = f"=SUM(D2:D{ws.max_row})"
ws.cell(row=ws.max_row + 1, column=4, value=sum_f)


doc = Document()
doc.add_heading('Отчёт', 0)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Имя'
hdr_cells[1].text = 'Количество'
hdr_cells[2].text = 'Цена'

for item in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item['name']
    row_cells[1].text = str(item['quantity'])
    row_cells[2].text = str(item['price'])

doc.add_page_break()  
doc.add_heading('Итоги', level=1)
total_sum = df['cost'].sum()
most_expensive_item = df.loc[df['price'].idxmax()]
doc.add_paragraph(f"Итоговая сумма: {total_sum:.2f} руб.")
doc.add_paragraph(f"Самая дорогая позиция: {most_expensive_item['name']}")

word_path = os.path.join(output_dir, 'report.docx')
doc.save(word_path)
print("Word создан")

pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)

pdf.cell(0, 10, txt="Report", ln=True, align='C')
pdf.ln(10)

col_width = pdf.w / 3.5
pdf.cell(col_width, 10, "Name", border=1)
pdf.cell(col_width, 10, "Quantity", border=1)
pdf.cell(col_width, 10, "Price", border=1)
pdf.ln()

for item in data:
    pdf.cell(col_width, 10, item['name'], border=1)
    pdf.cell(col_width, 10, str(item['quantity']), border=1)
    pdf.cell(col_width, 10, str(item['price']), border=1)
    pdf.ln()

pdf_path = os.path.join(output_dir, 'report.pdf')
pdf.output(pdf_path)
print("PDF создан")